from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import pandas as pd
pd.set_option("display.width", 10000)
pd.set_option("display.max_colwidth", None)
pd.set_option("display.max_columns", None)
import json

'''Funkce pro uvodni stranu'''
def vytvor_uvodni_stranu(dokument):
    # Vlastni text nadpisu strany
    styl = dokument.styles.add_style("vlastni_styl", 1) # 1 znaci, ze se jedna o styl odstavce
    font = styl.font
    font.name = "Courier New"
    font.size = Pt(64)
    font.underline = True

    # Pridani odstavce s nazvem dokumentu
    odstavec = dokument.add_paragraph("PŘEHLED PODEZŘELÝCH SUBJEKTŮ", style="vlastni_styl")
    odstavec.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Pridani obrazku s logem
    # Novy prazdny odstavec
    odstavec = dokument.add_paragraph()
    odstavec.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Pridani obsahu do existujiciho odstavce
    obsah_odstavce = odstavec.add_run()
    obsah_odstavce.add_picture("logo.jpg", width=Cm(4.5))

def pridej_stranku_se_subjektem(dokument, subjekt, tabulka_pro_report):
    # Pridani konce stranky
    dokument.add_page_break()

    # Nadpis s nazvem subjektu
    dokument.add_heading(f"Subjekt {subjekt}", level=0)

    ## Vlozeni grafu s rozdily
    # Nadpis nad grafy
    dokument.add_heading("Rozdíl mezi vykázanou a skutečnou produkcí", level=1)
    # Tabulka se dvema sloupci a jednim radkem pro vlozeni obou grafu vedle sebe
    tabulka = dokument.add_table(rows=1, cols=2) # Nova tabulka
    bunka_2023 = tabulka.cell(0, 0).paragraphs[0] # Pristup k prvni bunce a jejimu odstavci
    # Vlozeni noveho obsahu v podobe obrazku
    bunka_2023.add_run().add_picture(f"{subjekt}_2023.png", width=Cm(7.5))
    bunka_2024 = tabulka.cell(0, 1).paragraphs[0] # Pristup k prvni bunce a jejimu odstavci
    # Vlozeni noveho obsahu v podobe obrazku
    bunka_2024.add_run().add_picture(f"{subjekt}_2024.png", width=Cm(7.5))

    '''Tabulka s hodnotami produkce'''
    # Nadpis
    dokument.add_heading("Tabulka vykázaných produkcí", level=1)

    # Vytvoreni tabulky
    tabulka = dokument.add_table(rows=1, cols=tabulka_pro_report.shape[1])
    tabulka.style = "Light List"

    # Pridani textu hlavicky
    hlavicka = tabulka.rows[0].cells # "seznam" bunek prvniho radku
    for idx, sloupec in enumerate(tabulka_pro_report.columns):
        # Do i-te bunky radku se prida i-ty sloupec
        hlavicka[idx].text = sloupec

    # Iterace skrz radky tabulky
    for idx, radek in tabulka_pro_report.iterrows():
        # Pridani noveho radku tabulky
        radek_word = tabulka.add_row().cells
        # Iterace pres jednotlive hodnoty radku
        for idx_j, hodnota in enumerate(radek):
            # Prirazeni hodnoty bunce
            radek_word[idx_j].text = str(hodnota)
            # Zarovnani cisel v bunkach doprava
            radek_word[idx_j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Vytvoreni objektu Document
dokument = Document()
# Vytvoreni titulni strany
vytvor_uvodni_stranu(dokument)
# Nacteni dat o podezrelych subjektech
with open("problematicke_subjekty.json", "r") as f:
    problematicke_subjekty = json.load(f)
# Nacteni sloucenych dat
data = pd.read_csv("data_sloucena.csv")
# Iterace pres podezrele subjekty
for subjekt in problematicke_subjekty:
    # Filtrace tabulky podle subjektu
    tabulka_pro_report = data[data["id_entity"] == subjekt][["vyrobek_zkraceny", "2023_skut",
                                                             "2023_vyk", "2024_skut", "2024_vyk"]]
    # Pridani stranky se subjektu
    pridej_stranku_se_subjektem(dokument, subjekt, tabulka_pro_report)
# Ulozeni dokumentu
dokument.save("report_podezrelych_subjektu.docx")
